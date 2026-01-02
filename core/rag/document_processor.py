"""
Обработчик документов для RAG с использованием LangChain loaders и splitters
"""
import os
import asyncio
import io
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger
from config import settings

# LangChain imports
try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        Docx2txtLoader,
        UnstructuredExcelLoader,
        TextLoader,
    )
    # Пробуем импортировать HTML loader
    try:
        from langchain_community.document_loaders import BSHTMLLoader
        HTML_LOADER_AVAILABLE = True
    except ImportError:
        try:
            from langchain_community.document_loaders import UnstructuredHTMLLoader
            BSHTMLLoader = UnstructuredHTMLLoader
            HTML_LOADER_AVAILABLE = True
        except ImportError:
            HTML_LOADER_AVAILABLE = False
            BSHTMLLoader = None
    from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter
    from langchain_core.documents import Document as LangChainDocument
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    logger.warning(f"LangChain not fully available: {e}. Falling back to basic implementations.")
    LANGCHAIN_AVAILABLE = False
    LangChainDocument = None
    HTML_LOADER_AVAILABLE = False
    BSHTMLLoader = None

# Fallback imports для обратной совместимости (всегда импортируем, даже если LangChain доступен)
PyPDF2 = None
pypdf = None
Document = None
openpyxl = None

# Пробуем импортировать PyPDF2 (старая версия)
try:
    import PyPDF2
except (ImportError, NameError):
    PyPDF2 = None

# Пробуем импортировать pypdf (новая версия PyPDF2)
try:
    import pypdf
    # Если pypdf доступен, используем его как PyPDF2 для совместимости
    if PyPDF2 is None:
        PyPDF2 = pypdf
except (ImportError, NameError):
    pass

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

# Импорт для обработки HTML (fallback)
try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    BeautifulSoup = None

# Fallback на стандартную библиотеку html.parser
import html.parser

# Импорты для Vision API и конвертации PDF
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    Image = None

# Импорт Vision API клиента
try:
    from .vision_client import VisionAPIClient
    VISION_CLIENT_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Vision client not available: {e}")
    VISION_CLIENT_AVAILABLE = False
    VisionAPIClient = None

# Импорт LLM для очистки текста
try:
    from core.llm.factory import LLMProviderFactory
    from core.llm.base import LLMMessage
    from config import LLMProvider
    LLM_AVAILABLE = True
except ImportError as e:
    logger.warning(f"LLM not available for text cleaning: {e}")
    LLM_AVAILABLE = False
    LLMProviderFactory = None
    LLMMessage = None
    LLMProvider = None


class DocumentProcessor:
    """Класс для обработки различных типов документов с использованием LangChain"""
    
    def __init__(self, use_vision_api: bool = True, use_llm_cleaning: bool = None):
        """
        Инициализация процессора документов
        
        Args:
            use_vision_api: Использовать Google Vision API для извлечения текста (по умолчанию True)
            use_llm_cleaning: Использовать LLM для очистки текста (по умолчанию из settings)
        """
        # Инициализация text splitter из LangChain
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.rag_chunk_size,
                chunk_overlap=settings.rag_chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        else:
            self.text_splitter = None
        
        # Инициализация Vision API клиента
        self.use_vision_api = use_vision_api and VISION_CLIENT_AVAILABLE
        if self.use_vision_api:
            try:
                self.vision_client = VisionAPIClient()
                if not self.vision_client.is_available():
                    logger.warning("Vision API client initialized but API key is not set. Will use fallback methods.")
                    self.use_vision_api = False
                else:
                    logger.info("Vision API client initialized and ready to use")
            except Exception as e:
                logger.warning(f"Failed to initialize Vision API client: {e}. Will use fallback methods.")
                self.use_vision_api = False
        else:
            self.vision_client = None
        
        # Инициализация LLM для очистки текста
        self.use_llm_cleaning = use_llm_cleaning if use_llm_cleaning is not None else settings.use_llm_text_cleaning
        self.llm_provider = None
        if self.use_llm_cleaning and LLM_AVAILABLE:
            try:
                provider_type = settings.llm_text_cleaning_provider
                model = settings.llm_text_cleaning_model
                self.llm_provider = LLMProviderFactory.get_provider(provider_type, model)
                logger.info(f"LLM text cleaning enabled using {provider_type.value} with model {self.llm_provider.model}")
            except Exception as e:
                logger.warning(f"Failed to initialize LLM for text cleaning: {e}. Will use basic filtering.")
                self.use_llm_cleaning = False
    
    @staticmethod
    def _clean_ocr_text(text: str) -> str:
        """
        Очистка тексту від технічної інформації про PDF/OCR
        
        Видаляє:
        - Метадані PDF (Creator, Producer, CreationDate тощо)
        - Технічні рядки про розпізнавання
        - Артефакти OCR
        - Зайві пробіли та переноси рядків
        """
        if not text:
            return text
        
        # Видаляємо рядки з метаданими PDF
        lines = text.split('\n')
        cleaned_lines = []
        
        # Патерни для технічної інформації
        technical_patterns = [
            r'^PDF\s+Version',
            r'^Creator:',
            r'^Producer:',
            r'^CreationDate:',
            r'^ModDate:',
            r'^Title:',
            r'^Author:',
            r'^Subject:',
            r'^Keywords:',
            r'^Trapped:',
            r'^/Type\s+/',
            r'^/Subtype\s+/',
            r'^/Filter\s+/',
            r'^/Length\s+',
            r'^/Width\s+',
            r'^/Height\s+',
            r'^/ColorSpace\s+',
            r'^/BitsPerComponent\s+',
            r'^xref',
            r'^trailer',
            r'^startxref',
            r'^%%EOF',
            r'^/Page\s+',
            r'^/Pages\s+',
            r'^/MediaBox\s+',
            r'^/CropBox\s+',
            r'^/Rotate\s+',
            r'^/Parent\s+',
            r'^/Resources\s+',
            r'^/Font\s+',
            r'^/XObject\s+',
            r'^/ProcSet\s+',
            r'^/Contents\s+',
            r'OCR\s+confidence',
            r'Recognition\s+confidence',
            r'Text\s+extraction',
            r'Page\s+\d+\s+of\s+\d+',
            # Тестові рядки OCR
            r'This\s+is\s+text\s+in\s+English\s+for\s+OCR',
            r'for\s+OCR\s+recognition',
            r'OCR\s+recognition',
            r'Vision\s+API',
            r'Google\s+Vision\s+API',
            # Технічні описи документів
            r'PDF-файл[и]?\s+з\s+текстом',
            r'PDF-файл[и]?\s+с\s+текстом',
            r'PDF\s+файл[и]?\s+з\s+текстом',
            r'PDF\s+файл[и]?\s+с\s+текстом',
            r'для\s+распознавания\s+OCR',
            r'для\s+розпізнавання\s+OCR',
            r'документ[и]?\s+являются\s+PDF',
            r'документ[и]?\s+є\s+PDF',
            r'английском\s+языке\s+для\s+распознавания',
            r'англійською\s+мовою\s+для\s+розпізнавання',
            # Артефакти OCR (рядки з тільки крапками/символами)
            r'^[\.\s]+$',
            r'^[\-\s]+$',
            r'^[_\s]+$',
            r'^[=\s]+$',
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Перевіряємо, чи не є рядок технічною інформацією
            is_technical = False
            for pattern in technical_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    is_technical = True
                    break
            
            # Пропускаємо рядки, які виглядають як технічні команди PDF
            if line.startswith('/') and len(line) < 50:
                is_technical = True
            
            # Пропускаємо рядки з тільки цифрами та спецсимволами
            if re.match(r'^[\d\s/\\\[\](){}<>]+$', line) and len(line) < 30:
                is_technical = True
            
            # Пропускаємо дуже короткі рядки (менше 3 символів) - часто артефакти OCR
            if len(line) < 3:
                is_technical = True
            
            # Пропускаємо рядки, які містять тільки повторювані символи
            if len(line) > 0 and len(set(line.replace(' ', ''))) <= 2 and len(line) < 20:
                is_technical = True
            
            # Пропускаємо рядки з великою кількістю спецсимволів (більше 50%)
            if len(line) > 0:
                special_chars = len(re.findall(r'[^\w\s\u0400-\u04FF]', line))
                if special_chars / len(line) > 0.5:
                    is_technical = True
            
            if not is_technical:
                cleaned_lines.append(line)
        
        # Об'єднуємо рядки та очищаємо зайві пробіли
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Видаляємо множинні пробіли
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        # Видаляємо множинні переноси рядків (більше 2 підряд)
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    async def _clean_text_with_llm(self, text: str) -> str:
        """
        Очистка та структурування тексту через LLM (Ollama)
        
        Args:
            text: Текст для очищення
            
        Returns:
            Очищений та структурований текст
        """
        if not self.use_llm_cleaning or not self.llm_provider:
            return text
        
        if not text or not text.strip():
            return text
        
        try:
            # Промпт для очищення та структурування тексту
            system_prompt = """Ти - експерт з обробки документів. Твоя задача - очистити та структурувати текст, який був розпізнаний через OCR з PDF документів.

Ти маєш:
1. Видалити всю технічну інформацію про PDF (метадані, команди PDF, технічні рядки)
2. Видалити артефакти OCR (рядки з тільки символами, повторювані символи)
3. Видалити тестові рядки типу "This is text in English for OCR recognition"
4. Зберегти тільки корисний юридичний/бізнесовий зміст
5. Правильно структурувати текст (заголовки, параграфи, списки)
6. Виправити очевидні помилки розпізнавання
7. Зберегти оригінальну мову тексту (українська, російська, англійська)

Поверни тільки очищений та структурований текст, без додаткових пояснень."""

            user_prompt = f"""Очисти та структуруй наступний текст, який був розпізнаний через OCR:

{text[:8000]}"""  # Обмежуємо до 8000 символів для промпту
            
            messages = [
                LLMMessage(role="system", content=system_prompt),
                LLMMessage(role="user", content=user_prompt)
            ]
            
            logger.info(f"[DocumentProcessor] Cleaning text with LLM ({self.llm_provider.model})...")
            response = await self.llm_provider.generate(
                messages,
                temperature=0.3,  # Низька температура для більш точного результату
                max_tokens=10000
            )
            
            cleaned_text = response.content.strip()
            
            if cleaned_text and len(cleaned_text) > 0:
                logger.info(f"[DocumentProcessor] LLM cleaned text: {len(text)} -> {len(cleaned_text)} characters")
                return cleaned_text
            else:
                logger.warning("[DocumentProcessor] LLM returned empty text, using original")
                return text
                
        except Exception as e:
            logger.error(f"[DocumentProcessor] Error cleaning text with LLM: {e}")
            logger.debug(f"Falling back to basic filtering")
            # Fallback на базову фільтрацію
            return DocumentProcessor._clean_ocr_text(text)
    
    @staticmethod
    def _load_with_langchain(file_path: str) -> Optional[str]:
        """Загрузка документа с использованием LangChain loaders"""
        if not LANGCHAIN_AVAILABLE:
            return None
        
        try:
            file_ext = Path(file_path).suffix.lower()
            loader = None
            
            if file_ext == '.pdf':
                loader = PyPDFLoader(file_path)
            elif file_ext == '.docx':
                loader = Docx2txtLoader(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                loader = UnstructuredExcelLoader(file_path)
            elif file_ext == '.txt':
                loader = TextLoader(file_path, encoding='utf-8')
            elif file_ext in ['.html', '.htm'] and HTML_LOADER_AVAILABLE and BSHTMLLoader is not None:
                try:
                    loader = BSHTMLLoader(file_path)
                except Exception as e:
                    logger.warning(f"Failed to create BSHTMLLoader for {file_path}: {e}")
                    return None
            else:
                return None
            
            # Загрузка документов
            documents = loader.load()
            
            if not documents:
                logger.warning(f"LangChain loader returned no documents for {file_path}")
                return None
            
            # Объединение всех страниц/частей в один текст
            text = "\n\n".join([doc.page_content for doc in documents if doc.page_content])
            
            if not text or not text.strip():
                logger.warning(f"LangChain loader returned empty text for {file_path}")
                return None
            
            return text
            
        except Exception as e:
            logger.error(f"Error loading document with LangChain {file_path}: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def _pdf_to_images(self, file_path: str) -> List[bytes]:
        """
        Конвертация PDF в список изображений (байты)
        
        Args:
            file_path: Путь к PDF файлу
            
        Returns:
            Список байтов изображений (PNG)
        """
        if not PYMUPDF_AVAILABLE:
            logger.warning("PyMuPDF is not available. Cannot convert PDF to images.")
            return []
        
        try:
            images = []
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                # Конвертируем страницу в изображение с высоким разрешением
                mat = fitz.Matrix(2.0, 2.0)  # Увеличиваем разрешение в 2 раза
                pix = page.get_pixmap(matrix=mat)
                
                # Конвертируем в PNG байты
                img_bytes = pix.tobytes("png")
                images.append(img_bytes)
                logger.debug(f"Converted page {page_num + 1} to image: {len(img_bytes)} bytes")
            
            doc.close()
            logger.info(f"Converted PDF to {len(images)} images")
            return images
            
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return []
    
    def _extract_text_via_vision_api(self, file_path: str) -> Optional[str]:
        """
        Извлечение текста через Vision API (синхронная обертка)
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Извлеченный текст или None
        """
        if not self.use_vision_api or not self.vision_client:
            logger.debug(f"Vision API not available or client not initialized")
            return None
        
        logger.info(f"[DocumentProcessor] Starting Vision API extraction for file: {file_path}")
        
        try:
            # Запускаем асинхронный метод в синхронном контексте
            # В Celery tasks обычно нет event loop, поэтому используем asyncio.run()
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    # Если цикл уже запущен (редкий случай), используем ThreadPoolExecutor
                    logger.debug(f"[DocumentProcessor] Event loop is running, using ThreadPoolExecutor")
                    import concurrent.futures
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        future = executor.submit(
                            lambda: asyncio.run(self.vision_client.extract_text_from_file(file_path))
                        )
                        result = future.result(timeout=settings.vision_api_timeout + 10)
                        logger.info(f"[DocumentProcessor] Vision API extraction completed, result: {'success' if result is not None else 'failed'}")
                        return result
                else:
                    # Loop существует, но не запущен
                    logger.debug(f"[DocumentProcessor] Event loop exists but not running, using run_until_complete")
                    result = loop.run_until_complete(
                        self.vision_client.extract_text_from_file(file_path)
                    )
                    logger.info(f"[DocumentProcessor] Vision API extraction completed, result: {'success' if result is not None else 'failed'}")
                    return result
            except RuntimeError:
                # Нет event loop (обычный случай в Celery tasks)
                logger.debug(f"[DocumentProcessor] No event loop, using asyncio.run()")
                result = asyncio.run(
                    self.vision_client.extract_text_from_file(file_path)
                )
                logger.info(f"[DocumentProcessor] Vision API extraction completed, result: {'success' if result is not None else 'failed'}")
                return result
        except Exception as e:
            logger.error(f"[DocumentProcessor] Error extracting text via Vision API: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def _extract_text_from_pdf_via_vision(self, file_path: str) -> Optional[str]:
        """
        Извлечение текста из PDF через Vision API (конвертация страниц в изображения)
        
        Args:
            file_path: Путь к PDF файлу
            
        Returns:
            Извлеченный текст или None
        """
        if not self.use_vision_api or not self.vision_client:
            return None
        
        try:
            # Конвертируем PDF в изображения
            images = self._pdf_to_images(file_path)
            if not images:
                logger.warning("Failed to convert PDF to images")
                return None
            
            # Отправляем каждое изображение в Vision API и собираем текст
            all_text = []
            
            for i, image_bytes in enumerate(images):
                logger.info(f"[DocumentProcessor] Processing PDF page {i + 1}/{len(images)} via Vision API (size: {len(image_bytes)} bytes)")
                
                try:
                    # Запускаем асинхронный метод
                    # В Celery tasks нет event loop, поэтому всегда используем asyncio.run()
                    try:
                        loop = asyncio.get_event_loop()
                        if loop.is_running():
                            # Если loop уже запущен (редкий случай), используем ThreadPoolExecutor
                            logger.debug(f"[DocumentProcessor] Sending page {i + 1} image to Vision API (event loop running)")
                            import concurrent.futures
                            with concurrent.futures.ThreadPoolExecutor() as executor:
                                future = executor.submit(
                                    lambda: asyncio.run(
                                        self.vision_client.extract_text_from_image(
                                            image_bytes, 
                                            filename=f"page_{i+1}.png"
                                        )
                                    )
                                )
                                page_text = future.result(timeout=settings.vision_api_timeout + 10)
                        else:
                            # Loop существует, но не запущен
                            logger.debug(f"[DocumentProcessor] Sending page {i + 1} image to Vision API (using existing loop)")
                            page_text = loop.run_until_complete(
                                self.vision_client.extract_text_from_image(
                                    image_bytes,
                                    filename=f"page_{i+1}.png"
                                )
                            )
                    except RuntimeError:
                        # Нет event loop (обычный случай в Celery tasks)
                        logger.debug(f"[DocumentProcessor] Sending page {i + 1} image to Vision API (no event loop, using asyncio.run)")
                        page_text = asyncio.run(
                            self.vision_client.extract_text_from_image(
                                image_bytes,
                                filename=f"page_{i+1}.png"
                            )
                        )
                    
                    # page_text может быть None (ошибка) или строкой (включая пустую строку)
                    if page_text is not None:
                        if page_text.strip():
                            all_text.append(page_text)
                            logger.debug(f"Extracted {len(page_text)} characters from page {i + 1}")
                        else:
                            # Пустая строка - страница без текста, это валидный результат
                            logger.debug(f"Page {i + 1} contains no text (empty OCR result)")
                    else:
                        # None означает ошибку при обработке страницы
                        logger.warning(f"Failed to extract text from page {i + 1} via Vision API")
                        
                except Exception as page_error:
                    logger.warning(f"Error processing page {i + 1} via Vision API: {page_error}")
                    continue
            
            if all_text:
                combined_text = "\n\n".join(all_text)
                logger.info(f"Successfully extracted text from PDF via Vision API: {len(combined_text)} characters from {len(images)} pages")
                # Спочатку базова фільтрація
                cleaned_text = DocumentProcessor._clean_ocr_text(combined_text)
                if cleaned_text != combined_text:
                    logger.info(f"Basic filtering: removed {len(combined_text) - len(cleaned_text)} characters of technical metadata")
                
                # Потім очищення через LLM, якщо увімкнено
                if self.use_llm_cleaning and self.llm_provider:
                    try:
                        # Використовуємо синхронну обгортку для асинхронного методу
                        try:
                            loop = asyncio.get_event_loop()
                            if loop.is_running():
                                # Якщо loop вже запущений, використовуємо ThreadPoolExecutor
                                import concurrent.futures
                                with concurrent.futures.ThreadPoolExecutor() as executor:
                                    future = executor.submit(
                                        lambda: asyncio.run(self._clean_text_with_llm(cleaned_text))
                                    )
                                    cleaned_text = future.result(timeout=120)
                            else:
                                cleaned_text = loop.run_until_complete(self._clean_text_with_llm(cleaned_text))
                        except RuntimeError:
                            cleaned_text = asyncio.run(self._clean_text_with_llm(cleaned_text))
                    except Exception as e:
                        logger.warning(f"LLM cleaning failed: {e}, using basic filtered text")
                
                return cleaned_text
            else:
                logger.warning("No text extracted from PDF via Vision API")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF via Vision API: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF (с использованием Vision API и fallback на LangChain/PyPDF2)"""
        logger.info(f"Extracting text from PDF: {file_path}")
        
        # Сначала пробуем Vision API, если доступен
        if self.use_vision_api:
            logger.info(f"[DocumentProcessor] Attempting to extract text from PDF using Vision API: {file_path}")
            try:
                text = self._extract_text_from_pdf_via_vision(file_path)
                # Проверяем результат: None означает ошибку, пустая строка - валидный результат (нет текста)
                if text is not None:
                    if text.strip():
                        logger.info(f"[DocumentProcessor] Successfully extracted text from PDF using Vision API: {len(text)} characters")
                        return text
                    else:
                        # Пустая строка от Vision API - возможно, PDF не содержит текста
                        # Но попробуем fallback, так как Vision API мог пропустить текст
                        logger.warning(f"[DocumentProcessor] Vision API returned empty text for PDF (may not contain text or OCR failed), trying fallback methods")
                else:
                    # None означает ошибку при обработке
                    logger.warning(f"[DocumentProcessor] Vision API failed to process PDF, trying fallback methods")
            except Exception as e:
                logger.warning(f"[DocumentProcessor] Vision API failed for {file_path}: {e}, trying fallback methods")
                import traceback
                logger.debug(traceback.format_exc())
        
        # Пробуем LangChain loader
        if LANGCHAIN_AVAILABLE:
            logger.info("Attempting to extract text using LangChain PyPDFLoader...")
            try:
                text = DocumentProcessor._load_with_langchain(file_path)
                if text and text.strip():  # Проверяем, что текст не пустой
                    logger.info(f"Successfully extracted text from PDF using LangChain: {len(text)} characters")
                    return text
                else:
                    logger.warning(f"LangChain extracted empty text, trying PyPDF2 fallback")
            except Exception as e:
                logger.warning(f"LangChain PDF loader failed for {file_path}: {e}, trying fallback")
                import traceback
                logger.debug(traceback.format_exc())
        else:
            logger.warning("LangChain is not available, using PyPDF2 fallback")
        
        # Fallback на PyPDF2/pypdf (работает даже если LangChain доступен, но не смог извлечь текст)
        if PyPDF2 is None:
            logger.error("PyPDF2/pypdf is not installed. Cannot extract text from PDF.")
            logger.error("Please install PyPDF2 or pypdf: pip install PyPDF2 or pip install pypdf")
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                # Используем PyPDF2 или pypdf (они совместимы по API)
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                except AttributeError:
                    # Если PyPDF2 не имеет PdfReader, пробуем pypdf
                    if pypdf is not None:
                        pdf_reader = pypdf.PdfReader(file)
                    else:
                        raise
                
                num_pages = len(pdf_reader.pages)
                logger.debug(f"Extracting text from PDF with PyPDF2/pypdf: {num_pages} pages")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {i+1}: {page_error}")
                        continue
                
                if text.strip():
                    logger.debug(f"Successfully extracted text from PDF using PyPDF2/pypdf: {len(text)} characters")
                    return text
                else:
                    logger.warning(f"PyPDF2/pypdf extracted empty text from PDF: {file_path}")
                    return ""
        except NameError as e:
            # Обрабатываем случай, когда PyPDF2 не определен в области видимости
            logger.error(f"PyPDF2/pypdf is not available in this context: {e}")
            logger.error("This might be a Celery worker import issue. Please ensure PyPDF2 or pypdf is installed.")
            return ""
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path} with PyPDF2/pypdf: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX (с использованием Vision API и fallback на LangChain/python-docx)"""
        # Для DOCX сначала пробуем Vision API, если файл можно отправить как изображение
        # Но обычно DOCX лучше обрабатывать через python-docx, так что сначала пробуем стандартные методы
        
        # Пробуем LangChain loader
        if LANGCHAIN_AVAILABLE:
            try:
                text = DocumentProcessor._load_with_langchain(file_path)
                if text and text.strip():
                    logger.debug(f"Successfully extracted text from DOCX using LangChain: {len(text)} characters")
                    return text
            except Exception as e:
                logger.warning(f"LangChain DOCX loader failed for {file_path}: {e}, trying fallback")
        
        # Fallback на python-docx
        if Document is None:
            logger.error("python-docx is not installed. Cannot extract text from DOCX.")
            return ""
        
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            if text.strip():
                logger.debug(f"Successfully extracted text from DOCX using python-docx: {len(text)} characters")
                return text
            else:
                logger.warning(f"python-docx extracted empty text from DOCX: {file_path}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return ""
    
    def extract_text_from_xlsx(self, file_path: str) -> str:
        """Извлечение текста из XLSX (с fallback на openpyxl)"""
        # Пробуем LangChain loader
        if LANGCHAIN_AVAILABLE:
            try:
                text = DocumentProcessor._load_with_langchain(file_path)
                if text and text.strip():
                    logger.debug(f"Successfully extracted text from XLSX using LangChain: {len(text)} characters")
                    return text
            except Exception as e:
                logger.warning(f"LangChain XLSX loader failed for {file_path}: {e}, trying fallback")
        
        # Fallback на openpyxl
        if openpyxl is None:
            logger.error("openpyxl is not installed. Cannot extract text from XLSX.")
            return ""
        
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join([str(cell) if cell else "" for cell in row]) + "\n"
            if text.strip():
                logger.debug(f"Successfully extracted text from XLSX using openpyxl: {len(text)} characters")
                return text
            else:
                logger.warning(f"openpyxl extracted empty text from XLSX: {file_path}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from XLSX {file_path}: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Извлечение текста из TXT"""
        # Пробуем LangChain loader
        text = DocumentProcessor._load_with_langchain(file_path)
        if text is not None:
            return text
        
        # Fallback на старую реализацию
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return ""
    
    def _detect_html_encoding(self, file_path: str) -> str:
        """Определение кодировки HTML файла"""
        encodings = ['utf-8', 'windows-1251', 'cp1251', 'iso-8859-1', 'latin-1']
        
        # Сначала пытаемся найти кодировку в meta тегах
        try:
            with open(file_path, 'rb') as f:
                raw_content = f.read(8192)  # Читаем первые 8KB
                
                # Пробуем разные кодировки для чтения meta тегов
                for test_enc in ['utf-8', 'windows-1251', 'cp1251', 'iso-8859-1']:
                    try:
                        content_str = raw_content.decode(test_enc, errors='ignore')
                        
                        # Ищем charset в meta тегах
                        import re
                        charset_match = re.search(r'charset\s*=\s*["\']?([^"\'\s>]+)', content_str, re.IGNORECASE)
                        if charset_match:
                            detected_encoding = charset_match.group(1).lower()
                            # Нормализуем названия кодировок
                            if detected_encoding in ['windows-1251', 'cp1251']:
                                logger.debug(f"Detected encoding from meta tag: windows-1251")
                                return 'windows-1251'
                            elif detected_encoding in ['utf-8', 'utf8']:
                                logger.debug(f"Detected encoding from meta tag: utf-8")
                                return 'utf-8'
                            elif detected_encoding in ['iso-8859-1', 'latin-1']:
                                logger.debug(f"Detected encoding from meta tag: iso-8859-1")
                                return 'iso-8859-1'
                        break  # Если удалось декодировать, прекращаем попытки
                    except (UnicodeDecodeError, UnicodeError):
                        continue
        except Exception as e:
            logger.debug(f"Error detecting encoding from meta tags: {e}")
        
        # Если не нашли в meta, пробуем определить по содержимому
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='strict') as f:
                    f.read(1024)  # Пробуем прочитать первые 1KB
                logger.debug(f"Detected encoding by testing: {encoding}")
                return encoding
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # По умолчанию возвращаем utf-8
        logger.debug(f"Using default encoding: utf-8")
        return 'utf-8'
    
    def extract_text_from_html(self, file_path: str) -> str:
        """Извлечение текста из HTML (с использованием LangChain и fallback на BeautifulSoup/html.parser)"""
        logger.info(f"[DocumentProcessor] Processing HTML file: {file_path}")
        
        # Определяем кодировку файла
        encoding = self._detect_html_encoding(file_path)
        logger.debug(f"Detected HTML encoding: {encoding} for {file_path}")
        
        # Пробуем LangChain loader
        if LANGCHAIN_AVAILABLE and HTML_LOADER_AVAILABLE:
            try:
                text = DocumentProcessor._load_with_langchain(file_path)
                if text and text.strip():
                    logger.info(f"Successfully extracted text from HTML using LangChain: {len(text)} characters")
                    return text
            except Exception as e:
                logger.warning(f"LangChain HTML loader failed for {file_path}: {e}, trying fallback")
        
        # Fallback на BeautifulSoup с попыткой разных кодировок
        if BEAUTIFULSOUP_AVAILABLE:
            encodings_to_try = [encoding, 'utf-8', 'windows-1251', 'cp1251', 'iso-8859-1', 'latin-1']
            for enc in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=enc, errors='replace') as file:
                        html_content = file.read()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # Удаляем скрипты и стили
                        for script in soup(["script", "style"]):
                            script.decompose()
                        
                        # Извлекаем текст
                        text = soup.get_text()
                        
                        # Очищаем текст от лишних пробелов и переносов строк
                        lines = (line.strip() for line in text.splitlines())
                        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                        text = '\n'.join(chunk for chunk in chunks if chunk)
                        
                        if text.strip():
                            logger.info(f"Successfully extracted text from HTML using BeautifulSoup (encoding: {enc}): {len(text)} characters")
                            return text
                except UnicodeDecodeError:
                    logger.debug(f"Failed to decode HTML with encoding {enc}, trying next")
                    continue
                except Exception as e:
                    logger.warning(f"BeautifulSoup failed for {file_path} with encoding {enc}: {e}, trying next")
                    continue
            
            logger.warning(f"BeautifulSoup failed for {file_path} with all encodings, trying html.parser fallback")
        
        # Fallback на стандартную библиотеку html.parser с попыткой разных кодировок
        encodings_to_try = [encoding, 'utf-8', 'windows-1251', 'cp1251', 'iso-8859-1', 'latin-1']
        for enc in encodings_to_try:
            try:
                class HTMLTextExtractor(html.parser.HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                        self.skip_tags = {'script', 'style', 'meta', 'link', 'head'}
                        self.in_skip_tag = False
                    
                    def handle_starttag(self, tag, attrs):
                        if tag in self.skip_tags:
                            self.in_skip_tag = True
                    
                    def handle_endtag(self, tag):
                        if tag in self.skip_tags:
                            self.in_skip_tag = False
                        elif tag in {'p', 'div', 'br', 'li'}:
                            self.text.append('\n')
                    
                    def handle_data(self, data):
                        if not self.in_skip_tag:
                            self.text.append(data.strip())
                
                with open(file_path, 'r', encoding=enc, errors='replace') as file:
                    html_content = file.read()
                    extractor = HTMLTextExtractor()
                    extractor.feed(html_content)
                    
                    # Объединяем текст и очищаем от лишних пробелов
                    text = ' '.join(extractor.text)
                    text = re.sub(r'\s+', ' ', text)  # Заменяем множественные пробелы на один
                    text = re.sub(r'\n\s*\n', '\n', text)  # Удаляем пустые строки
                    
                    if text.strip():
                        logger.info(f"Successfully extracted text from HTML using html.parser (encoding: {enc}): {len(text)} characters")
                        return text.strip()
            except UnicodeDecodeError:
                logger.debug(f"Failed to decode HTML with encoding {enc}, trying next")
                continue
            except Exception as e:
                logger.warning(f"html.parser failed for {file_path} with encoding {enc}: {e}, trying next")
                continue
        
        logger.error(f"Failed to extract text from HTML {file_path} with all methods and encodings")
        return ""
    
    def process_document(self, file_path: str) -> str:
        """Обработка документа любого поддерживаемого типа"""
        file_ext = Path(file_path).suffix.lower()
        
        # Проверяем, является ли файл изображением - если да, отправляем напрямую в Vision API
        if self.use_vision_api and file_ext.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp']:
            logger.info(f"[DocumentProcessor] Detected image file: {file_path}, will send to Vision API")
            text = self._extract_text_via_vision_api(file_path)
            # text может быть None (ошибка) или пустой строкой (нет текста на изображении)
            if text is not None:
                # Очищаем текст от технической информации
                cleaned_text = DocumentProcessor._clean_ocr_text(text)
                if cleaned_text != text:
                    logger.info(f"[DocumentProcessor] Cleaned OCR text: removed {len(text) - len(cleaned_text)} characters of technical metadata")
                # Пустая строка - валидный результат (изображение без текста)
                logger.info(f"[DocumentProcessor] Vision API processing completed for image: {file_path}")
                return cleaned_text
            else:
                logger.warning(f"[DocumentProcessor] Vision API failed to process image: {file_path}, no fallback available for images")
                return ""
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.xlsx' or file_ext == '.xls':
            return self.extract_text_from_xlsx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_ext in ['.html', '.htm']:
            return self.extract_text_from_html(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
        """Разбиение текста на чанки с использованием LangChain text splitter"""
        # Используем LangChain text splitter если доступен
        if self.text_splitter and LANGCHAIN_AVAILABLE:
            # Обновляем параметры если переданы
            if chunk_size is not None or chunk_overlap is not None:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size or settings.rag_chunk_size,
                    chunk_overlap=chunk_overlap or settings.rag_chunk_overlap,
                    length_function=len,
                    separators=["\n\n", "\n", " ", ""]
                )
            else:
                splitter = self.text_splitter
            
            # Создаем временный документ для splitter
            doc = LangChainDocument(page_content=text)
            chunks = splitter.split_documents([doc])
            return [chunk.page_content for chunk in chunks]
        
        # Fallback на старую реализацию
        chunk_size = chunk_size or settings.rag_chunk_size
        chunk_overlap = chunk_overlap or settings.rag_chunk_overlap
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - chunk_overlap
        
        return chunks

